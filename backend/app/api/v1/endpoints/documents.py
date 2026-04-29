from __future__ import annotations

import json
import tempfile
import time
import uuid
import logging
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Dict, Any

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import get_current_user
from app.db.session import get_db
from app.models.document import Document, DocumentStatus
from app.models.extraction import Extraction
from app.models.user import User
from app.models.workflow import Workflow
from app.schemas.document import (
    ActionLogEntry,
    BatchProcessRequest,
    BatchProcessResponse,
    DocumentResultsResponse,
    DocumentStatusResponse,
    DocumentUploadResponse,
)

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/documents", tags=["documents"])

ALLOWED_CONTENT_TYPES = [
    "application/pdf",
    "image/png",
    "image/jpeg",
    "image/tiff",
    "text/plain",
    "text/markdown",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
]


@router.post("/upload", response_model=DocumentUploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    # Accept workflow_id from BOTH form-data (preferred) and query string (legacy).
    # Without Form(...) FastAPI treats it as a query param only, so frontends
    # that send it in multipart silently lose the workflow association.
    workflow_id: Optional[str] = Form(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Upload a document file."""
    if file.content_type not in ALLOWED_CONTENT_TYPES:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.content_type}")

    doc_id = uuid.uuid4()
    file_content = await file.read()
    remote_path = f"uploads/{doc_id}/{file.filename}"

    # Store file via MinIO (best-effort; fallback to local file)
    try:
        from app.services.storage import get_storage
        storage = get_storage()
        storage.upload_bytes(file_content, remote_path, content_type=file.content_type)
    except Exception as e:
        logger.warning(f"Storage upload failed, saving locally: {e}")

    # Always save a local copy as fallback for OCR processing
    local_dir = Path("uploads") / str(doc_id)
    local_dir.mkdir(parents=True, exist_ok=True)
    local_path = local_dir / file.filename
    local_path.write_bytes(file_content)
    logger.info(f"Saved local copy: {local_path}")

    document = Document(
        id=doc_id,
        filename=f"{doc_id}_{file.filename}",
        original_filename=file.filename,
        content_type=file.content_type,
        file_size=len(file_content),
        storage_path=remote_path,
        status=DocumentStatus.UPLOADED,
        workflow_id=uuid.UUID(workflow_id) if workflow_id else None,
        uploaded_by=current_user.id,
    )
    db.add(document)
    await db.commit()
    await db.refresh(document)

    return DocumentUploadResponse(
        document_id=str(document.id),
        file_name=file.filename,
        upload_timestamp=document.created_at,
    )


@router.post("/process/{document_id}")
async def process_document(
    document_id: str,
    workflow_id: str = "generic",
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Process a document: OCR → LLM extraction → save results."""
    result = await db.execute(select(Document).where(Document.id == document_id))
    document = result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    document.status = DocumentStatus.PROCESSING
    await db.commit()

    start_time = time.time()

    try:
        # --- Step 1: Download file from storage & OCR ---
        ocr_text = ""
        ocr_engine_used = "tesseract"
        tmp_path = None

        # Try to download from MinIO storage
        try:
            from app.services.storage import get_storage
            storage = get_storage()
            with tempfile.NamedTemporaryFile(suffix=Path(document.original_filename).suffix, delete=False) as tmp:
                storage.download(document.storage_path, Path(tmp.name))
                tmp_path = Path(tmp.name)
            logger.info(f"Downloaded from storage: {document.storage_path}")
        except Exception as e:
            logger.warning(f"Storage download failed: {e}")

            # Fallback: look for file locally (saved during upload)
            local_candidates = [
                Path(document.storage_path) if document.storage_path else None,
                Path(f"uploads/{document.storage_path}") if document.storage_path else None,
                # Direct path: uploads/<doc_id>/<filename>
                Path(f"uploads/{document_id}/{document.original_filename}"),
            ]
            for lp in local_candidates:
                if lp and lp.exists():
                    tmp_path = lp
                    logger.info(f"Found file locally: {lp}")
                    break

        # Run OCR / text extraction if we have a file
        if tmp_path and tmp_path.exists():
            file_ext = tmp_path.suffix.lower()

            # Fast path: plain text / markdown — read directly, no OCR needed
            if file_ext in (".txt", ".md"):
                try:
                    ocr_text = tmp_path.read_text(encoding="utf-8", errors="replace")
                    ocr_engine_used = "text-passthrough"
                    logger.info(f"Text passthrough for {document_id}: {len(ocr_text)} chars")
                except Exception as e:
                    logger.warning(f"Text read failed: {e}")

            # Fast path: .docx — extract via python-docx
            elif file_ext == ".docx":
                try:
                    from docx import Document as DocxDocument
                    docx_doc = DocxDocument(str(tmp_path))
                    ocr_text = "\n".join(p.text for p in docx_doc.paragraphs if p.text)
                    ocr_engine_used = "docx-passthrough"
                    logger.info(f"DOCX passthrough for {document_id}: {len(ocr_text)} chars")
                except Exception as e:
                    logger.warning(f"DOCX read failed: {e}, falling back to OCR")

            # OCR path: PDF + images
            if not ocr_text:
                try:
                    from app.config import settings as app_settings
                    if app_settings.OCR_PROVIDER == "mistral" and app_settings.MISTRAL_API_KEY:
                        from app.services.ocr import MistralOCREngine, TesseractOCREngine, OCRPipeline
                        primary = MistralOCREngine(api_key=app_settings.MISTRAL_API_KEY)
                        fallback = TesseractOCREngine()
                        ocr_engine_used = "mistral"
                    else:
                        from app.services.ocr import TesseractOCREngine, OCRPipeline
                        primary = TesseractOCREngine()
                        fallback = None
                        ocr_engine_used = "tesseract"

                    pipeline = OCRPipeline(primary, fallback)
                    ocr_results = pipeline.process_document(tmp_path)
                    ocr_text = "\n\n".join(r.text for r in ocr_results if r.text)
                    logger.info(f"OCR completed for {document_id}: {len(ocr_results)} pages, {len(ocr_text)} chars of text")
                except Exception as e:
                    logger.warning(f"OCR processing failed: {e}")

        # Fallback: use pre-existing ocr_text from document (e.g., seeded data)
        if not ocr_text and document.ocr_text:
            ocr_text = document.ocr_text
            ocr_engine_used = "pre-stored"
            logger.info(f"Using pre-stored OCR text for {document_id}: {len(ocr_text)} chars")

        # --- Step 3: Get workflow extraction schema ---
        field_defs = []
        # Auto-detect workflow from filename if none was selected
        if not document.workflow_id:
            fname_lower = document.original_filename.lower()
            if any(kw in fname_lower for kw in ("cv", "resume", "curriculum")):
                wf_result = await db.execute(
                    select(Workflow).where(Workflow.name.ilike("%cv%skill%"))
                )
                auto_wf = wf_result.scalar_one_or_none()
                if auto_wf:
                    document.workflow_id = auto_wf.id
                    logger.info(f"Auto-detected CV workflow for {document_id}: {auto_wf.id}")

        wf = None
        extraction_engine = None
        if document.workflow_id:
            wf_result = await db.execute(select(Workflow).where(Workflow.id == document.workflow_id))
            wf = wf_result.scalar_one_or_none()
            if wf and wf.extraction_schema:
                field_defs = wf.extraction_schema.get("fields", [])
                # Workflow-level engine override (e.g. granite-docling for forms)
                extraction_engine = wf.extraction_schema.get("extraction_engine")

        # ── Granite-Docling multimodal path (forms with signatures / handwriting) ──
        if extraction_engine == "granite-docling" and tmp_path and tmp_path.exists():
            return await _process_with_granite_docling(
                document=document,
                tmp_path=tmp_path,
                wf=wf,
                field_defs=field_defs or [],
                ocr_text=ocr_text,
                start_time=start_time,
                db=db,
            )

        if not field_defs:
            # Fallback: guess from filename
            fname = document.original_filename.lower()
            if "inv" in fname:
                field_defs = [
                    {"name": "vendor_name", "label": "Vendor Name", "field_type": "text"},
                    {"name": "invoice_number", "label": "Invoice Number", "field_type": "text"},
                    {"name": "invoice_date", "label": "Invoice Date", "field_type": "date"},
                    {"name": "due_date", "label": "Due Date", "field_type": "date"},
                    {"name": "subtotal", "label": "Subtotal", "field_type": "currency"},
                    {"name": "tax_amount", "label": "Tax Amount", "field_type": "currency"},
                    {"name": "total_amount", "label": "Total Amount", "field_type": "currency"},
                    {"name": "payment_terms", "label": "Payment Terms", "field_type": "text"},
                    {"name": "po_number", "label": "PO Number", "field_type": "text"},
                ]
            elif "receipt" in fname:
                field_defs = [
                    {"name": "merchant_name", "label": "Merchant", "field_type": "text"},
                    {"name": "receipt_date", "label": "Date", "field_type": "date"},
                    {"name": "total_amount", "label": "Total", "field_type": "currency"},
                    {"name": "tax_amount", "label": "Tax", "field_type": "currency"},
                    {"name": "payment_method", "label": "Payment Method", "field_type": "text"},
                    {"name": "category", "label": "Category", "field_type": "text"},
                ]
            elif any(kw in fname for kw in ("cv", "resume", "curriculum")):
                field_defs = [
                    {"name": "candidate_name", "label": "Candidate Name", "field_type": "text"},
                    {"name": "email", "label": "Email", "field_type": "text"},
                    {"name": "phone", "label": "Phone", "field_type": "text"},
                    {"name": "location", "label": "Location", "field_type": "text"},
                    {"name": "professional_summary", "label": "Professional Summary", "field_type": "text"},
                    {"name": "technical_skills", "label": "Technical Skills", "field_type": "text"},
                    {"name": "years_of_experience", "label": "Years of Experience", "field_type": "number"},
                    {"name": "current_role", "label": "Current Role/Title", "field_type": "text"},
                    {"name": "current_company", "label": "Current Company", "field_type": "text"},
                    {"name": "work_history", "label": "Work History", "field_type": "text"},
                    {"name": "education", "label": "Education", "field_type": "text"},
                    {"name": "certifications", "label": "Certifications", "field_type": "text"},
                    {"name": "languages", "label": "Languages", "field_type": "text"},
                    {"name": "expected_salary", "label": "Expected Salary", "field_type": "text"},
                ]
            else:
                field_defs = [
                    {"name": "title", "label": "Title", "field_type": "text"},
                    {"name": "date", "label": "Date", "field_type": "date"},
                    {"name": "author", "label": "Author", "field_type": "text"},
                    {"name": "summary", "label": "Summary", "field_type": "text"},
                ]

        # --- Step 4: LLM extraction ---
        field_names = [f["name"] for f in field_defs]
        field_desc = "\n".join(f"- {f['name']}: {f.get('label', f['name'])} ({f.get('field_type', 'text')})" for f in field_defs)

        prompt_text = f"""Extract the following fields from this document. Return a JSON object with field names as keys and objects with "value" (string) and "confidence" (float 0-1) as values.

Fields to extract:
{field_desc}

Document text:
---
{ocr_text if ocr_text else f"[No OCR text available. Document filename: {document.original_filename}]"}
---

Return ONLY valid JSON. Example format:
{{"vendor_name": {{"value": "ACME Corp", "confidence": 0.95}}, "total_amount": {{"value": "$1,234.56", "confidence": 0.92}}}}"""

        from app.services.llm_provider import LLMProviderFactory
        llm = LLMProviderFactory.from_settings()
        # Scale max_tokens with field count: ~120 tokens per field for {value, confidence}
        # pairs with descriptions. Floor at 4096, cap at 16000.
        max_out_tokens = max(4096, min(len(field_defs) * 150 + 500, 16000))
        llm_response = await llm.chat(
            messages=[
                {"role": "system", "content": "You are a document extraction assistant. Extract structured fields from document text. Always return valid JSON only, no markdown."},
                {"role": "user", "content": prompt_text},
            ],
            temperature=0.0,
            max_tokens=max_out_tokens,
        )

        # Parse LLM response
        response_text = llm_response.content.strip()
        # Strip markdown code fences if present
        if response_text.startswith("```"):
            response_text = response_text.split("\n", 1)[-1]
            if response_text.endswith("```"):
                response_text = response_text[:-3]
            response_text = response_text.strip()

        extracted_fields = json.loads(response_text)
        logger.info(f"LLM extracted {len(extracted_fields)} fields for {document_id}: {list(extracted_fields.keys())}")

        # --- Step 5: Save extraction results ---
        # Clear any previous extractions for this document (supports re-processing)
        from sqlalchemy import delete as sa_delete
        await db.execute(sa_delete(Extraction).where(Extraction.document_id == document.id))

        processing_time = (time.time() - start_time) * 1000
        for field_name, field_data in extracted_fields.items():
            if field_name not in field_names:
                continue
            raw_value = field_data.get("value", "") if isinstance(field_data, dict) else field_data
            # Convert list values to comma-separated string for DB storage
            if isinstance(raw_value, list):
                value = ", ".join(str(v) for v in raw_value)
            else:
                value = str(raw_value) if raw_value is not None else ""
            confidence = field_data.get("confidence", 0.85) if isinstance(field_data, dict) else 0.85

            extraction = Extraction(
                id=uuid.uuid4(),
                document_id=document.id,
                field_name=field_name,
                field_value=value,
                field_type=next((f.get("field_type", "text") for f in field_defs if f["name"] == field_name), "text"),
                confidence=confidence,
                model_used=llm_response.model,
            )
            db.add(extraction)

        document.status = DocumentStatus.COMPLETED
        document.metadata_json = {
            "ocr_engine": ocr_engine_used,
            "processing_time_ms": processing_time,
            "model_used": llm_response.model,
            "fields_extracted": len(extracted_fields),
        }
        await db.commit()

        logger.info(f"Document processed: {document_id}, {len(extracted_fields)} fields in {processing_time:.0f}ms")
        return {"status": "completed", "document_id": document_id, "fields_extracted": len(extracted_fields), "processing_time_ms": round(processing_time)}

    except Exception as e:
        logger.error(f"Document processing failed: {document_id} - {e}")
        document.status = DocumentStatus.FAILED
        document.metadata_json = {"error": str(e)}
        await db.commit()
        return {"status": "failed", "document_id": document_id, "error": str(e)}


# ── Granite-Docling multimodal extraction path ────────────────────────────────
#
# Used when a workflow has extraction_engine="granite-docling" in its schema.
# Captures: full text, handwritten field crops, and signature thumbnails.
# Persists signatures + handwriting samples in document.metadata_json so the
# review UI can render them for audit.

async def _process_with_granite_docling(
    document: Document,
    tmp_path: Path,
    wf: Optional[Workflow],
    field_defs: List[Dict[str, Any]],
    ocr_text: str,
    start_time: float,
    db: AsyncSession,
) -> Dict[str, Any]:
    from app.services.granite_docling_engine import GraniteDoclingEngine, get_engine

    if not GraniteDoclingEngine.is_available():
        document.status = DocumentStatus.FAILED
        document.metadata_json = {
            "error": "granite-docling engine not installed on backend",
            "install_hint": "pip install -r backend/requirements-granite.txt",
        }
        await db.commit()
        return {"status": "failed", "document_id": str(document.id), "error": "granite-docling not installed"}

    logger.info(f"[granite-docling] processing {document.id}")
    engine = get_engine()
    granite_result = engine.process(tmp_path)

    if granite_result.status == "failed":
        document.status = DocumentStatus.FAILED
        document.metadata_json = {
            "ocr_engine": "granite-docling",
            "error": granite_result.error_message or "granite-docling processing failed",
        }
        await db.commit()
        return {"status": "failed", "document_id": str(document.id), "error": granite_result.error_message}

    # ---- Map granite output → workflow schema fields via the LLM -----------
    # The LLM only fills the printed-text-derived fields. Signature presence
    # and handwriting count come directly from the granite multimodal output
    # (no LLM hallucination risk).

    field_names = [f["name"] for f in field_defs]
    field_desc = "\n".join(
        f"- {f['name']}: {f.get('description') or f.get('label') or f['name']} ({f.get('type', 'text')})"
        for f in field_defs
        # Skip the structural fields we'll fill ourselves
        if f["name"] not in {"guard_signature_present", "supervisor_signature_present", "handwritten_field_count"}
    )

    # Compose form context: markdown body + handwritten fields list
    form_text_blob = granite_result.markdown or granite_result.plain_text or ocr_text or ""
    handwritten_fields = [f for f in granite_result.form_fields if f.is_handwritten]
    handwritten_summary = "\n".join(
        f"- {h.label or '(no label)'}: {h.value or '(no value)'}"
        for h in handwritten_fields[:30]
    ) or "(no handwritten fields detected)"

    prompt_text = f"""Extract the requested fields from this filled security guard attendance form.
The form was processed by IBM granite-docling-258M, which separately identified
handwritten content from printed text.

Fields to extract:
{field_desc}

Document body (markdown, layout preserved):
---
{form_text_blob[:6000]}
---

Handwritten fields detected (label: value):
{handwritten_summary}

Return ONLY valid JSON: {{"field_name": {{"value": "...", "confidence": 0.0-1.0}}, ...}}.
If a field is not in the document, return {{"value": null, "confidence": 0.0}}."""

    extracted_fields: Dict[str, Any] = {}
    llm_model_used: Optional[str] = None
    try:
        from app.services.llm_provider import LLMProviderFactory
        llm = LLMProviderFactory.from_settings()
        max_out = max(2048, min(len(field_defs) * 150 + 500, 8000))
        llm_response = await llm.chat(
            messages=[
                {"role": "system", "content": "You extract structured fields from forms. Return valid JSON only."},
                {"role": "user", "content": prompt_text},
            ],
            temperature=0.0,
            max_tokens=max_out,
        )
        llm_model_used = llm_response.model
        response_text = llm_response.content.strip()
        if response_text.startswith("```"):
            response_text = response_text.split("\n", 1)[-1]
            if response_text.endswith("```"):
                response_text = response_text[:-3]
            response_text = response_text.strip()
        extracted_fields = json.loads(response_text)
    except Exception as e:
        logger.warning(f"[granite-docling] LLM mapping failed, falling back to handwritten-only: {e}")
        # Fallback: just copy handwritten fields into the schema where labels match
        lower_field_names = {n.lower() for n in field_names}
        for h in handwritten_fields:
            label = (h.label or "").lower().replace(" ", "_")
            if label in lower_field_names and h.value:
                extracted_fields[label] = {"value": h.value, "confidence": 0.6}

    # ---- Inject the structural fields from granite output (authoritative) -
    sig_count = len(granite_result.signatures)
    hw_count = len(handwritten_fields)
    extracted_fields["guard_signature_present"] = {
        "value": "true" if sig_count >= 1 else "false",
        "confidence": 0.95,
    }
    extracted_fields["supervisor_signature_present"] = {
        "value": "true" if sig_count >= 2 else "false",
        "confidence": 0.85,
    }
    extracted_fields["handwritten_field_count"] = {
        "value": str(hw_count),
        "confidence": 1.0,
    }

    # ---- Save extractions --------------------------------------------------
    from sqlalchemy import delete as sa_delete
    await db.execute(sa_delete(Extraction).where(Extraction.document_id == document.id))

    for f in field_defs:
        name = f["name"]
        data = extracted_fields.get(name) or extracted_fields.get(name.lower()) or {}
        raw_value = data.get("value") if isinstance(data, dict) else data
        if isinstance(raw_value, list):
            value = ", ".join(str(v) for v in raw_value)
        elif raw_value is None:
            value = ""
        else:
            value = str(raw_value)
        confidence = float(data.get("confidence", 0.0)) if isinstance(data, dict) else 0.0
        db.add(Extraction(
            id=uuid.uuid4(),
            document_id=document.id,
            field_name=name,
            field_value=value,
            field_type=f.get("type", "text"),
            confidence=confidence,
            model_used=f"granite-docling+{llm_model_used or 'fallback'}",
        ))

    # ---- Run validation against rules + decide review_status --------------
    rules = (wf.validation_rules or {}).get("rules", []) if wf and wf.validation_rules else []
    validation_results: List[Dict[str, Any]] = []
    rejected_reasons: List[str] = []

    def _field_value(name: str) -> str:
        d = extracted_fields.get(name) or extracted_fields.get(name.lower()) or {}
        v = d.get("value") if isinstance(d, dict) else d
        return "" if v is None else str(v).strip()

    for rule in rules:
        rule_type = rule.get("type")
        passed = True
        msg = ""
        try:
            if rule_type == "boolean_true":
                fv = _field_value(rule.get("field", "")).lower()
                passed = fv in ("true", "yes", "1")
                msg = f"{rule.get('field')} = {fv or 'empty'}"
            elif rule_type == "min_count":
                fv_str = _field_value(rule.get("field", ""))
                try:
                    fv_num = int(fv_str) if fv_str else 0
                except ValueError:
                    fv_num = 0
                min_required = int(rule.get("min", 1))
                passed = fv_num >= min_required
                msg = f"{rule.get('field')} = {fv_num} (need ≥ {min_required})"
            elif rule_type == "not_empty":
                fv = _field_value(rule.get("field", ""))
                passed = bool(fv)
                msg = f"{rule.get('field')} = {fv or '(empty)'}"
            else:
                # Unknown rule type — soft-pass so we don't block the form
                passed = True
                msg = f"unsupported rule type: {rule_type}"
        except Exception as e:
            passed = False
            msg = f"rule evaluation error: {e}"

        validation_results.append({
            "rule": rule.get("rule"),
            "description": rule.get("description"),
            "passed": passed,
            "severity": rule.get("severity", "warning"),
            "detail": msg,
        })
        if not passed and rule.get("severity") == "rejection":
            rejected_reasons.append(rule.get("description") or rule.get("rule") or "validation failed")

    review_status = "rejected" if rejected_reasons else "approved"
    processing_time = (time.time() - start_time) * 1000

    # ---- Persist visual artifacts to metadata_json ------------------------
    document.status = DocumentStatus.COMPLETED
    document.metadata_json = {
        "ocr_engine": "granite-docling",
        "model_used": llm_model_used or "granite-docling",
        "processing_time_ms": processing_time,
        "fields_extracted": len([k for k in extracted_fields if k in field_names]),
        "review_status": review_status,
        "rejected_reasons": rejected_reasons,
        "validation_results": validation_results,
        "granite": {
            "page_count": granite_result.page_count,
            "signature_count": sig_count,
            "handwritten_field_count": hw_count,
            "headings": granite_result.headings,
            "signatures": [
                {
                    "page_number": s.page_number,
                    "bbox": s.bbox,
                    "confidence": s.confidence,
                    "thumbnail_base64": s.thumbnail_base64,
                }
                for s in granite_result.signatures
            ],
            "handwritten_fields": [
                {
                    "page_number": h.page_number,
                    "label": h.label,
                    "value": h.value,
                    "bbox": h.bbox,
                    "thumbnail_base64": h.thumbnail_base64,
                }
                for h in handwritten_fields
            ],
            # Full markdown extraction so reviewers can read the form text
            "markdown": form_text_blob[:8000],
        },
    }
    await db.commit()

    logger.info(
        f"[granite-docling] done {document.id}: review_status={review_status}, "
        f"sigs={sig_count}, handwritten={hw_count}, time={processing_time:.0f}ms"
    )
    return {
        "status": "completed",
        "document_id": str(document.id),
        "review_status": review_status,
        "rejected_reasons": rejected_reasons,
        "signature_count": sig_count,
        "handwritten_field_count": hw_count,
        "processing_time_ms": round(processing_time),
    }


@router.post("/process-batch", response_model=BatchProcessResponse)
async def process_batch(
    request: BatchProcessRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Process multiple documents in batch."""
    batch_id = str(uuid.uuid4())

    for doc_id in request.document_ids:
        result = await db.execute(select(Document).where(Document.id == doc_id))
        document = result.scalar_one_or_none()
        if document:
            document.status = DocumentStatus.PROCESSING
    await db.commit()

    logger.info(f"Batch processing queued: {batch_id}, {len(request.document_ids)} documents")

    return BatchProcessResponse(
        batch_id=batch_id,
        document_count=len(request.document_ids),
        created_at=datetime.utcnow(),
    )


@router.get("/status/{document_id}", response_model=DocumentStatusResponse)
async def get_document_status(
    document_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Get processing status of a document."""
    result = await db.execute(select(Document).where(Document.id == document_id))
    document = result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    progress = {
        DocumentStatus.UPLOADED: 0,
        DocumentStatus.PROCESSING: 50,
        DocumentStatus.EXTRACTED: 75,
        DocumentStatus.VALIDATED: 90,
        DocumentStatus.COMPLETED: 100,
        DocumentStatus.FAILED: 0,
    }

    return DocumentStatusResponse(
        document_id=str(document.id),
        status=document.status.value,
        progress_percent=progress.get(document.status, 0),
        created_at=document.created_at,
        updated_at=document.updated_at,
    )


@router.get("/results/{document_id}", response_model=DocumentResultsResponse)
async def get_document_results(
    document_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Get extraction results for a completed document."""
    result = await db.execute(select(Document).where(Document.id == document_id))
    document = result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    # Build results from extractions
    fields = {}
    for ext in document.extractions:
        fields[ext.field_name] = {
            "value": ext.field_value,
            "confidence": ext.confidence,
            "raw_text": ext.corrected_value or ext.field_value,
        }

    # Extract error message from metadata_json
    error_message = None
    metadata = document.metadata_json or {}
    if isinstance(metadata, dict):
        error_message = metadata.get("error")

    # Build action log entries
    action_log_entries = []
    for log in document.action_logs:
        action_log_entries.append(ActionLogEntry(
            action_type=log.action_type,
            status=log.status,
            action_config=log.action_config,
            result=log.result,
            error_message=log.error_message,
            created_at=log.created_at,
        ))

    md = document.metadata_json or {}
    md_engine = md.get("ocr_engine") if isinstance(md, dict) else None
    return DocumentResultsResponse(
        document_id=str(document.id),
        file_name=document.original_filename,
        status=document.status.value,
        fields=fields,
        tables={},
        processing_time_ms=float(md.get("processing_time_ms") or 0.0),
        ocr_engine=md_engine or "mistral",
        completed_at=document.updated_at,
        error_message=error_message,
        action_logs=action_log_entries,
        metadata=md if isinstance(md, dict) else None,
    )


@router.get("/list", response_model=List[DocumentStatusResponse])
async def list_documents(
    workflow_id: Optional[str] = None,
    status: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """List documents, optionally filtered."""
    query = select(Document).where(Document.uploaded_by == current_user.id)

    if workflow_id:
        query = query.where(Document.workflow_id == workflow_id)
    if status:
        query = query.where(Document.status == DocumentStatus(status))

    result = await db.execute(query.order_by(Document.created_at.desc()))
    documents = result.scalars().all()

    progress_map = {
        DocumentStatus.UPLOADED: 0,
        DocumentStatus.PROCESSING: 50,
        DocumentStatus.EXTRACTED: 75,
        DocumentStatus.VALIDATED: 90,
        DocumentStatus.COMPLETED: 100,
        DocumentStatus.FAILED: 0,
    }

    return [
        DocumentStatusResponse(
            document_id=str(doc.id),
            file_name=doc.original_filename,
            status=doc.status.value,
            progress_percent=progress_map.get(doc.status, 0),
            created_at=doc.created_at,
            updated_at=doc.updated_at,
        )
        for doc in documents
    ]


# ---------- Skill-Fit Assessment ----------

def _load_role_skill_matrix() -> dict:
    """Load role-skill-matrix.json from templates."""
    matrix_paths = [
        Path(__file__).resolve().parents[4] / "templates" / "hr" / "data" / "role-skill-matrix.json",
        Path("templates/hr/data/role-skill-matrix.json"),
        Path("../templates/hr/data/role-skill-matrix.json"),
    ]
    for p in matrix_paths:
        if p.exists():
            return json.loads(p.read_text())
    return {"roles": {}}


def _fuzzy_match(candidate_item: str, required_item: str) -> bool:
    """Case-insensitive matching for skills/certs with synonym awareness."""
    c = candidate_item.lower().strip()
    r = required_item.lower().strip()
    if r in c or c in r:
        return True

    # Synonym groups — if the required skill is a category, match against known members
    _SYNONYMS = {
        "programming": ["python", "javascript", "typescript", "java", "go", "c++", "c#",
                         "ruby", "rust", "dart", "kotlin", "swift", "php", "scala", "r"],
        "data structures": ["algorithms", "data structures & algorithms", "dsa"],
        "cloud services": ["aws", "gcp", "azure", "cloud", "ec2", "s3", "lambda",
                           "cloudfront", "vertex ai", "cloud run", "sagemaker"],
        "ci/cd": ["github actions", "gitlab ci", "jenkins", "argocd", "ci/cd pipeline",
                  "continuous integration", "continuous deployment"],
        "unit testing": ["testing", "unit test", "integration test", "code coverage", "jest", "pytest"],
        "machine learning": ["ml", "deep learning", "neural network", "tensorflow", "pytorch",
                             "scikit-learn", "xgboost", "lightgbm", "catboost"],
        "data analysis": ["data analytics", "analytics", "exploratory data analysis", "eda",
                          "pandas", "numpy", "data visualization"],
        "data visualization": ["matplotlib", "seaborn", "plotly", "tableau", "power bi", "looker",
                               "streamlit", "grafana"],
        "risk assessment": ["credit risk", "market risk", "operational risk", "risk management",
                            "risk analysis", "risk modeling"],
        "regulatory knowledge": ["ojk", "pojk", "pbi", "basel", "ifrs", "regulatory compliance"],
        "excel advanced": ["excel", "vba", "solver", "excel advanced"],
    }
    for category, members in _SYNONYMS.items():
        if r == category and any(m in c for m in members):
            return True
        if r in members and c == category:
            return True
    return False


def _match_skills(candidate_skills: List[str], required: List[str]) -> tuple:
    """Returns (matched, missing) lists."""
    matched = []
    missing = []
    for req in required:
        if any(_fuzzy_match(cs, req) for cs in candidate_skills):
            matched.append(req)
        else:
            missing.append(req)
    return matched, missing


def _compute_fit_score(
    extractions: dict,
    role: dict,
) -> dict:
    """Compute weighted skill-fit score for a candidate against a role."""
    # Parse candidate data from extractions
    def get_val(key: str) -> str:
        field = extractions.get(key, {})
        return field.get("value", "") if isinstance(field, dict) else str(field)

    def get_list(key: str) -> List[str]:
        val = get_val(key)
        if not val:
            return []
        return [s.strip() for s in val.split(",") if s.strip()]

    candidate_skills = get_list("technical_skills") + get_list("soft_skills")
    candidate_certs = get_list("certifications")
    candidate_langs = get_list("languages_spoken")

    try:
        candidate_exp = float(get_val("total_years_experience") or "0")
    except (ValueError, TypeError):
        candidate_exp = 0.0

    candidate_education = get_val("education_level").lower()
    candidate_major = get_val("education_major").lower()

    # --- Score factors (weights from DEMO-SCENARIO.md) ---
    scores = {}

    # 1. Required Skills Match (35%)
    req_skills = role.get("required_skills", [])
    prog_req = role.get("programming_languages", {}).get("required", [])
    all_required = req_skills + prog_req
    matched_req, missing_req = _match_skills(candidate_skills, all_required)
    req_ratio = len(matched_req) / max(len(all_required), 1)
    scores["required_skills"] = {"score": round(req_ratio * 100), "weight": 35, "matched": matched_req, "missing": missing_req}

    # 2. Experience (20%)
    min_exp = role.get("min_experience_years", 0)
    if min_exp == 0:
        exp_score = 100
    elif candidate_exp >= min_exp * 1.5:
        exp_score = 100
    elif candidate_exp >= min_exp:
        exp_score = 80 + 20 * (candidate_exp - min_exp) / max(min_exp * 0.5, 1)
    elif candidate_exp > 0:
        exp_score = (candidate_exp / min_exp) * 70
    else:
        exp_score = 0
    scores["experience"] = {"score": round(min(exp_score, 100)), "weight": 20,
                            "candidate": candidate_exp, "required": min_exp}

    # 3. Preferred Skills Match (15%)
    pref_skills = role.get("preferred_skills", [])
    prog_pref = role.get("programming_languages", {}).get("preferred", [])
    all_preferred = pref_skills + prog_pref
    matched_pref, missing_pref = _match_skills(candidate_skills, all_preferred)
    pref_ratio = len(matched_pref) / max(len(all_preferred), 1)
    scores["preferred_skills"] = {"score": round(pref_ratio * 100), "weight": 15, "matched": matched_pref, "missing": missing_pref}

    # 4. Education Level (10%)
    edu_map = {"s3": 100, "s2": 90, "s1": 75, "d4": 65, "d3": 55}
    edu_score = 50  # default
    for key, val in edu_map.items():
        if key in candidate_education:
            edu_score = val
            break
    # Bonus for matching major
    preferred_majors = [m.lower() for m in role.get("preferred_majors", [])]
    major_match = any(pm in candidate_major or candidate_major in pm for pm in preferred_majors) if candidate_major else False
    if major_match:
        edu_score = min(edu_score + 10, 100)
    scores["education"] = {"score": edu_score, "weight": 10, "level": get_val("education_level"),
                           "major": get_val("education_major"), "major_match": major_match}

    # 5. Certifications (10%)
    req_certs = role.get("certifications", {}).get("required", [])
    pref_certs = role.get("certifications", {}).get("preferred", [])
    matched_req_cert, missing_req_cert = _match_skills(candidate_certs, req_certs)
    matched_pref_cert, _ = _match_skills(candidate_certs, pref_certs)
    if req_certs:
        cert_score = (len(matched_req_cert) / len(req_certs)) * 70 + (len(matched_pref_cert) / max(len(pref_certs), 1)) * 30
    else:
        cert_score = (len(matched_pref_cert) / max(len(pref_certs), 1)) * 100 if pref_certs else 50
    scores["certifications"] = {"score": round(min(cert_score, 100)), "weight": 10,
                                "matched_required": matched_req_cert, "missing_required": missing_req_cert,
                                "matched_preferred": matched_pref_cert, "candidate_certs": candidate_certs}

    # 6. Industry Relevance (10%) — based on overall profile match
    industry_score = 50  # neutral baseline
    if candidate_exp >= min_exp:
        industry_score += 20
    if major_match:
        industry_score += 15
    if matched_req_cert:
        industry_score += 15
    scores["industry_relevance"] = {"score": min(industry_score, 100), "weight": 10}

    # --- Weighted total ---
    total = sum(s["score"] * s["weight"] for s in scores.values()) / 100.0
    total = round(total, 1)

    # Fit level
    if total >= 85:
        fit_level = "Strong Fit"
        fit_color = "green"
        recommendation = "Proceed to interview"
    elif total >= 70:
        fit_level = "Good Fit"
        fit_color = "blue"
        recommendation = "Consider for interview, minor gaps"
    elif total >= 50:
        fit_level = "Partial Fit"
        fit_color = "yellow"
        recommendation = "Significant gaps, assess if trainable"
    else:
        fit_level = "Weak Fit"
        fit_color = "red"
        recommendation = "Does not meet minimum requirements"

    return {
        "overall_score": total,
        "fit_level": fit_level,
        "fit_color": fit_color,
        "recommendation": recommendation,
        "factors": scores,
        "candidate_name": get_val("candidate_name"),
        "target_role": role.get("title", ""),
        "salary_band": role.get("salary_band", {}),
    }


@router.get("/skill-fit/{document_id}")
async def get_skill_fit_assessment(
    document_id: str,
    role: str = "software_engineer",
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Assess candidate-role fit by comparing extracted CV data against role-skill-matrix."""
    result = await db.execute(select(Document).where(Document.id == document_id))
    document = result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    # Build extractions dict
    fields = {}
    for ext in document.extractions:
        fields[ext.field_name] = {"value": ext.field_value, "confidence": ext.confidence}

    if not fields:
        raise HTTPException(status_code=400, detail="No extraction data available. Process the document first.")

    # Load role matrix
    matrix = _load_role_skill_matrix()
    available_roles = {k: v.get("title", k) for k, v in matrix.get("roles", {}).items()}

    role_data = matrix.get("roles", {}).get(role)
    if not role_data:
        raise HTTPException(status_code=400, detail=f"Unknown role: {role}. Available: {list(available_roles.keys())}")

    assessment = _compute_fit_score(fields, role_data)
    assessment["available_roles"] = available_roles

    return assessment


@router.get("/roles")
async def list_available_roles(
    current_user: User = Depends(get_current_user),
):
    """List all available roles from the skill matrix."""
    matrix = _load_role_skill_matrix()
    roles = []
    for key, val in matrix.get("roles", {}).items():
        roles.append({
            "id": key,
            "title": val.get("title", key),
            "department": val.get("department", ""),
            "min_experience": val.get("min_experience_years", 0),
        })
    return {"roles": roles}
