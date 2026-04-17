from __future__ import annotations

from pathlib import Path
from typing import Optional, List
from dataclasses import dataclass
from .ocr import OCRPipeline, OCRResult, OCRStatus
import logging

logger = logging.getLogger(__name__)


@dataclass
class ProcessedDocument:
    """Container for a processed document with all extracted pages."""
    document_id: str
    file_path: Path
    file_type: str
    total_pages: int
    ocr_results: List[OCRResult]
    status: str  # "success", "partial", "failed"
    error_message: Optional[str] = None


class DocumentProcessor:
    """Orchestrates document processing workflow."""

    def __init__(self, ocr_pipeline: OCRPipeline):
        self.ocr_pipeline = ocr_pipeline

    def process(self, document_path: Path, document_id: str) -> ProcessedDocument:
        try:
            file_type = document_path.suffix.lower()

            if file_type in [".png", ".jpg", ".jpeg", ".tiff"]:
                preprocessed = self.ocr_pipeline.preprocess_image(document_path)
                ocr_results = [self.ocr_pipeline.process_page(preprocessed)]
            elif file_type in [".txt", ".md"]:
                # Plain-text documents: read directly, no OCR needed
                text = document_path.read_text(encoding="utf-8", errors="replace")
                ocr_results = [OCRResult(
                    page_number=1, text=text, confidence=1.0,
                    status=OCRStatus.SUCCESS, engine="text-passthrough",
                )]
            elif file_type in [".doc", ".docx"]:
                # Word documents: extract text via python-docx (.docx only; .doc not supported)
                text = ""
                try:
                    from docx import Document as DocxDocument
                    docx = DocxDocument(str(document_path))
                    text = "\n".join(p.text for p in docx.paragraphs if p.text)
                except Exception as docx_err:
                    logger.warning(f"docx extraction failed, falling back to OCR: {docx_err}")
                    ocr_results = self.ocr_pipeline.process_document(document_path)
                else:
                    ocr_results = [OCRResult(
                        page_number=1, text=text, confidence=1.0,
                        status=OCRStatus.SUCCESS, engine="docx-passthrough",
                    )]
            else:
                ocr_results = self.ocr_pipeline.process_document(document_path)

            status = "success" if ocr_results and ocr_results[0].text else "partial"

            return ProcessedDocument(
                document_id=document_id,
                file_path=document_path,
                file_type=file_type,
                total_pages=len(ocr_results),
                ocr_results=ocr_results,
                status=status,
            )
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return ProcessedDocument(
                document_id=document_id,
                file_path=document_path,
                file_type=document_path.suffix.lower(),
                total_pages=0,
                ocr_results=[],
                status="failed",
                error_message=str(e),
            )
